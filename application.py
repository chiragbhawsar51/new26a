from flask import Flask, render_template, request, send_file
from docxtpl import DocxTemplate
from docx import Document
import datetime
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx2pdf import convert  # To convert DOCX to PDF
import os  # For removing temporary files
import mammoth

app = Flask(__name__, static_url_path='/static')
application = app

# Constants
COVER_LETTER_TEMPLATE = "Cover_letterr.docx"
FINAL_FILE_DOCX_FILENAME = "Final_Cover_letter_with_table_{}.docx"
FINAL_FILE_PDF_FILENAME = "Final_Cover_letter_with_table_{}.pdf"

app = Flask(__name__)

def generate_cover_letter(context):
    today_date = datetime.datetime.today().strftime('%B %d, %Y')
    context['today_date'] = today_date

    doc = DocxTemplate(COVER_LETTER_TEMPLATE)
    doc.render(context)
    
    temp_filename = "Temp_Cover_letter.docx"
    doc.save(temp_filename)

    return temp_filename

def create_and_insert_table(doc, target_index, records):
    num_cols = 5  # Assuming 5 columns for S.no, Description, Rate, Quantity, and Amount
    table = doc.add_table(rows=len(records) + 1, cols=num_cols)  # Add one for header row

    headers = ["S.no", "Description", "Rate", "Quantity", "Amount"]
    for i, header_text in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header_text
        shading_color = "808080"  # Gray header
        cell._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}" w:val="clear"/>'))
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # White header text
        cell.paragraphs[0].runs[0].font.size = Pt(12)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in table.rows:
        row.cells[0].width = Inches(0.3)
        row.cells[1].width = Inches(3.0)
        row.cells[2].width = Inches(0.5)
        row.cells[3].width = Inches(0.5)
        row.cells[4].width = Inches(0.5)

    for i, record in enumerate(records, start=1):
        for j, header_text in enumerate(headers):
            cell = table.cell(i, j)
            if header_text == "Amount":
                rate = float(record[2])
                quantity = float(record[3])
                amount = rate * quantity
                cell.text = str(amount)
            else:
                cell.text = str(record[j])
            
            shading_color = "D3D3D3" if i % 2 == 0 else "FFFFFF"  # Gray row, white row text
            cell._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}" w:val="clear"/>'))
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Black text
            cell.paragraphs[0].runs[0].font.size = Pt(12)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    table.style = 'Table Grid'
    paragraph = doc.paragraphs[target_index]
    paragraph.insert_paragraph_before()._p.addnext(table._tbl)

def convert_docx_to_html(file_path):
    with open(file_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value
        return html

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        try:
            context = {
                'offer_name': request.form['offer_name'],
                'offer_no': request.form['offer_no'],
                'company_name': request.form['company_name'],
                'city_name': request.form['city_name'],
                'state_name': request.form['state_name'],
                'manager_name': request.form['manager_name'],
                'enquiry_sub': request.form['enquiry_sub'],
                'delivery_dates': request.form['delivery_dates'],
                'your_name': request.form['your_name'],
                'contact_no': request.form['contact_no']
            }
            records_count = int(request.form['records_count'])
            records = []
            for i in range(records_count):
                sn = request.form[f'sn_{i}']
                description = request.form[f'description_{i}']
                rate = float(request.form[f'rate_{i}'])
                quantity = float(request.form[f'quantity_{i}'])
                records.append((sn, description, rate, quantity))

            cover_letter_file = generate_cover_letter(context)
            doc = Document(cover_letter_file)

            target_text = "Annexure II-Commercial Terms and Conditions."
            target_index = None
            for i, paragraph in enumerate(doc.paragraphs):
                if target_text in paragraph.text:
                    target_index = i
                    break

            if target_index is not None:
                create_and_insert_table(doc, target_index, records)
                today_date = datetime.datetime.today().strftime('%Y%m%d')
                final_docx_file = FINAL_FILE_DOCX_FILENAME.format(today_date)
                doc.save(final_docx_file)

                final_pdf_file = FINAL_FILE_PDF_FILENAME.format(today_date)
                convert(final_docx_file, final_pdf_file)

                html_content = convert_docx_to_html(final_docx_file)
                os.remove(final_docx_file)

                return render_template('preview.html', html_content=html_content, filename=final_pdf_file)
            else:
                return "Error: Target paragraph not found in the document."
        except Exception as e:
            return str(e)
    return render_template('index.html')

@app.route('/download/<filename>')
def download(filename):
    return send_file(filename, as_attachment=True)

if __name__ == '__main__':
    app.run('0.0.0.0')

